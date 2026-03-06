"""
FORENSIC ANALYSIS WEB APPLICATION
Modern, Futuristic Frontend with Backend Analysis
Supports: RAW, DD, E01 disk images and USB drives
Exports: PDF and DOC formats
"""

from flask import Flask, render_template, request, jsonify, send_file, session
from flask_cors import CORS
from werkzeug.utils import secure_filename
import os
import json
import threading
from datetime import datetime
from pathlib import Path
from forensic_pipeline_advanced import AdvancedForensicAnalyzer, HAS_PYTSK3, HAS_PYEWF
import re
# NOTE: reportlab + docx are imported lazily inside export functions to speed up startup
# NOTE: rag_engine is imported lazily below (on first RAG use) to avoid slow startup time

app = Flask(__name__)
app.secret_key = 'forensic-analysis-secret-key-2026'
CORS(app, 
     resources={r"/api/*": {"origins": ["http://localhost:5173", "http://localhost:3000", "http://127.0.0.1:5173", "http://127.0.0.1:3000"]}},
     supports_credentials=True,
     methods=["GET", "POST", "OPTIONS"],
     allow_headers=["Content-Type"],
     expose_headers=["Content-Disposition", "Content-Type", "Content-Length"])

# Configuration
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
OUTPUT_FOLDER = os.path.join(BASE_DIR, 'analysis_output')
ALLOWED_EXTENSIONS = {'raw', 'dd', 'e01', 'img', 'iso', 'bin', 'dmg'}

# Create folders if they don't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Global analysis status tracker
analysis_status = {
    'current': None,
    'progress': 0,
    'status': 'idle',
    'message': '',
    'errors': [],
    'results': None,
    'rag_collection': None
}

def allowed_file(filename):
    """Check if file extension is allowed."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def run_forensic_analysis(image_path, analysis_id):
    """Run forensic analysis in background thread."""
    global analysis_status
    
    try:
        analysis_status['current'] = analysis_id
        analysis_status['status'] = 'analyzing'
        analysis_status['progress'] = 5
        analysis_status['message'] = 'Initializing forensic engine...'

        # Auto-detect whether real analysis libraries are available
        use_demo = not (HAS_PYTSK3 and HAS_PYEWF)
        if use_demo:
            analysis_status['message'] = (
                'Note: pytsk3/pyewf not installed — running enriched demo mode '
                'with real file metadata.'
            )
            print('[!] pytsk3/pyewf not available — falling back to demo mode')

        analyzer = AdvancedForensicAnalyzer(image_path, use_demo=use_demo)

        # Enrich demo data with real file information from the uploaded image
        if use_demo and os.path.exists(image_path):
            real_name = os.path.basename(image_path)
            real_size = os.path.getsize(image_path)
            import hashlib as _hashlib
            sha = _hashlib.sha256()
            with open(image_path, 'rb') as _f:
                for chunk in iter(lambda: _f.read(65536), b''):
                    sha.update(chunk)
            real_hash = sha.hexdigest()
            # Inject real image info as the first file entry
            real_key = f'0_{real_name}'
            analyzer.files_metadata[real_key] = {
                'name': real_name,
                'path': f'/{real_name}',
                'size': real_size,
                'inode': 0,
                'creation_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'modification_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'access_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'is_deleted': False,
                'is_directory': False,
            }
            analyzer.file_hashes[real_key] = real_hash
            from pathlib import Path as _Path
            ext = _Path(real_name).suffix.lower() or '(no extension)'
            analyzer.file_extension_stats[ext] += 1
            print(f'[+] Real file: {real_name} | Size: {real_size:,} bytes | SHA256: {real_hash[:16]}...')
        
        analysis_status['progress'] = 25
        analysis_status['message'] = 'Opening disk image...'
        if not analyzer.open_disk_image():
            raise Exception("Failed to open disk image")
        
        analysis_status['progress'] = 35
        analysis_status['message'] = 'Detecting partitions...'
        if not analyzer.detect_partitions():
            raise Exception("Failed to detect partitions")
        
        analysis_status['progress'] = 45
        analysis_status['message'] = 'Opening filesystem...'
        if not analyzer.open_filesystem():
            raise Exception("Failed to open filesystem")
        
        analysis_status['progress'] = 55
        analysis_status['message'] = 'Scanning files recursively...'
        if not analyzer.recursively_scan_files():
            raise Exception("Failed to scan files")
        
        analysis_status['progress'] = 65
        analysis_status['message'] = 'Computing file hashes...'
        if not analyzer.compute_file_hashes():
            raise Exception("Failed to compute hashes")
        
        analysis_status['progress'] = 75
        analysis_status['message'] = 'Detecting encryption and network artifacts...'
        encrypted_items = analyzer.detect_encryption()
        network_artifacts = analyzer.detect_network_artifacts()
        
        analysis_status['progress'] = 85
        analysis_status['message'] = 'Generating reports...'
        
        # Generate JSON report
        json_path = os.path.join(OUTPUT_FOLDER, f'{analysis_id}_report.json')
        if not analyzer.export_to_json(json_path):
            raise Exception("Failed to export JSON")
        
        # Generate timeline report
        timeline_path = os.path.join(OUTPUT_FOLDER, f'{analysis_id}_timeline.txt')
        if not analyzer.generate_timeline_report(timeline_path):
            raise Exception("Failed to generate timeline")
        
        # Generate extension report
        extension_path = os.path.join(OUTPUT_FOLDER, f'{analysis_id}_extensions.txt')
        if not analyzer.generate_extension_report(extension_path):
            raise Exception("Failed to generate extension report")
        
        # Generate summary report
        summary_path = os.path.join(OUTPUT_FOLDER, f'{analysis_id}_summary.txt')
        if not analyzer.generate_summary_report(json_path, summary_path):
            raise Exception("Failed to generate summary")
        
        analysis_status['progress'] = 90
        analysis_status['message'] = 'Finalizing analysis...'
        
        # Prepare results data
        with open(json_path, 'r', encoding='utf-8') as f:
            json_data = json.load(f)
        
        # ── RAG: chunk + vectorize the report into ChromaDB ────────────
        analysis_status['progress'] = 95
        analysis_status['message'] = 'Vectorizing report for RAG chatbot...'
        try:
            _ensure_rag_loaded()  # Lazy-load RAG on first use
            collection_name = _vectorize_and_store(analysis_id, json_data)
            analysis_status['rag_collection'] = collection_name
            print(f'[+] RAG collection ready: {collection_name}')
        except Exception as vec_err:
            print(f'[!] RAG vectorization warning: {vec_err}')
            analysis_status['rag_collection'] = None

        analysis_status['results'] = {
            'json_path': json_path,
            'summary_path': summary_path,
            'timeline_path': timeline_path,
            'extension_path': extension_path,
            'analysis_id': analysis_id,
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'image_path': image_path,
            'data': json_data,
            'rag_collection': analysis_status['rag_collection']
        }
        
        analysis_status['progress'] = 100
        analysis_status['status'] = 'completed'
        analysis_status['message'] = 'Analysis completed successfully!'
        
    except Exception as e:
        import traceback as _tb
        analysis_status['status'] = 'error'
        analysis_status['message'] = f'Error: {str(e)}'
        analysis_status['errors'].append(str(e))
        print(f"[-] Forensic Analysis Error: {str(e)}", flush=True)
        _tb.print_exc()

def _llm_generate_section(prompt_text, fallback=""):
    """Ask Llama 3 via Ollama to generate a report section. Returns fallback on failure."""
    import requests as _req
    try:
        resp = _req.post(
            "http://localhost:11434/api/generate",
            json={"model": "llama3", "prompt": prompt_text, "stream": False},
            timeout=120,
        )
        if resp.status_code == 200:
            return resp.json().get("response", "").strip() or fallback
    except Exception as exc:
        print(f"[!] LLM report section error: {exc}")
    return fallback


def _build_report_data_summary(data):
    """Build a concise, fact-only data summary string for LLM prompts."""
    summary = data.get('summary', {})
    files = data.get('files', {})
    encrypted = data.get('encrypted_items', [])
    network = data.get('network_artifacts', [])
    extensions = data.get('file_extension_statistics', {})
    partition = data.get('partition_info', {})

    suspicious_files = [
        f for f in files.values()
        if any(f['name'].lower().endswith(ext) for ext in ['.exe', '.bat', '.ps1', '.dll', '.scr', '.vbs'])
    ]
    deleted_files = [f for f in files.values() if f.get('is_deleted')]

    lines = [
        f"Image path: {data.get('image_path', 'N/A')}",
        f"Scan timestamp: {data.get('scan_timestamp', 'N/A')}",
        f"Filesystem type: {partition.get('filesystem_type', 'Unknown')}",
        f"Block size: {partition.get('block_size', 'N/A')} bytes",
        f"Block count: {partition.get('block_count', 'N/A')}",
        f"Total files scanned: {summary.get('total_files', 0)}",
        f"Total deleted files found: {summary.get('total_deleted_files', 0)}",
        f"Suspicious files count (by extension): {summary.get('suspicious_files_count', len(suspicious_files))}",
        f"Encrypted items count: {summary.get('encrypted_items_count', len(encrypted))}",
        f"Network artifacts count: {summary.get('network_artifacts_count', len(network))}",
        f"File types detected: {len(extensions)}",
    ]
    # Top extensions
    top_ext = sorted(extensions.items(), key=lambda x: x[1], reverse=True)[:10]
    if top_ext:
        lines.append("Top file extensions: " + ", ".join(f"{e} ({c})" for e, c in top_ext))
    # Suspicious file names
    if suspicious_files:
        lines.append("Suspicious file names: " + ", ".join(f['name'] for f in suspicious_files[:15]))
    else:
        lines.append("Suspicious files: NONE found.")
    # Deleted file names
    if deleted_files:
        lines.append("Deleted file names (sample): " + ", ".join(
            f['name'] for f in sorted(deleted_files, key=lambda x: x['modification_time'], reverse=True)[:10]))
    else:
        lines.append("Deleted files: NONE found.")
    # Encrypted
    if encrypted:
        lines.append("Encrypted item names: " + ", ".join(i['name'] for i in encrypted[:10]))
    else:
        lines.append("Encrypted items: NONE found.")
    # Network
    if network:
        lines.append("Network artifact names: " + ", ".join(i['name'] for i in network[:10]))
    else:
        lines.append("Network artifacts: NONE found.")
    return "\n".join(lines)


def _pdf_page_footer(canvas_obj, doc_obj):
    """Draw page number and thin header line on every page (except cover)."""
    page_num = canvas_obj.getPageNumber()
    if page_num == 1:  # skip cover page
        return
    canvas_obj.saveState()
    # ── thin top rule ──
    canvas_obj.setStrokeColor(colors.HexColor('#1a237e'))
    canvas_obj.setLineWidth(0.5)
    canvas_obj.line(0.6*inch, letter[1] - 0.45*inch, letter[0] - 0.6*inch, letter[1] - 0.45*inch)
    canvas_obj.setFont('Helvetica', 7)
    canvas_obj.setFillColor(colors.HexColor('#1a237e'))
    canvas_obj.drawString(0.6*inch, letter[1] - 0.40*inch, "FORENSIC ANALYSIS REPORT")
    canvas_obj.drawRightString(letter[0] - 0.6*inch, letter[1] - 0.40*inch, "CONFIDENTIAL")
    # ── bottom page number ──
    canvas_obj.setFont('Helvetica', 8)
    canvas_obj.setFillColor(colors.HexColor('#888888'))
    canvas_obj.drawCentredString(letter[0] / 2, 0.35*inch, f"Page {page_num}")
    canvas_obj.restoreState()


def _fmt_size(size_bytes):
    """Format byte size into human-readable string."""
    if size_bytes >= 1073741824:
        return f"{size_bytes / (1024**3):.2f} GB"
    if size_bytes >= 1048576:
        return f"{size_bytes / (1024**2):.2f} MB"
    if size_bytes >= 1024:
        return f"{size_bytes / 1024:.1f} KB"
    return f"{size_bytes} B"


def generate_pdf_report(analysis_id, results):
    """Generate a professional PDF forensic report with fact-only LLM analysis."""
    try:
        # Lazy import — only loaded when user actually requests PDF export
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
        from reportlab.lib import colors

        pdf_path = os.path.join(OUTPUT_FOLDER, f'{analysis_id}_report.pdf')
        data = results['data']['forensic_report']
        summary = data['summary']
        data_summary_text = _build_report_data_summary(data)
        image_path = data.get('image_path', 'N/A')
        scan_ts = data.get('scan_timestamp', 'N/A')

        # ── Ask LLM for executive summary — STRICT fact-only prompt ───
        exec_summary = _llm_generate_section(
            f"""You are a digital forensics report writer.

STRICT RULES — you MUST follow every one:
1. ONLY use the facts provided in the DATA section below. Do NOT assume, speculate, or invent any information.
2. If a count is zero, state clearly that none were found. Do NOT fabricate threats or risks.
3. Do NOT reference external tools, external scans, or anything not in the data.
4. Do NOT use markdown formatting. Write plain text only.
5. Write in formal, professional forensic report language.

Write an EXECUTIVE SUMMARY (3-5 short paragraphs) covering ONLY:
- What disk image was analyzed (use the exact image path from the data)
- Filesystem type and scan date from the data
- Exact counts: total files, deleted files, suspicious files, encrypted items, network artifacts
- A factual statement of whether items of interest were found or not

=== DATA (facts only) ===
{data_summary_text}
=== END DATA ===""",
            fallback="Executive summary could not be generated by the AI model. Please review the statistical tables below for all findings."
        )

        # ── Ask LLM for findings narrative — STRICT fact-only ─────────
        findings_narrative = _llm_generate_section(
            f"""You are a digital forensics report writer.

STRICT RULES:
1. ONLY reference facts from the DATA section below. Do NOT assume, speculate, or add information not in the data.
2. If a category has zero items, state that clearly and move on.
3. Do NOT use markdown. Write plain text only.
4. Write in formal forensic report language.

Write a DETAILED FINDINGS section (4-6 paragraphs) covering:
- Suspicious files: list each by name if present, or state none were found
- Deleted files: list names if present, or state none were found
- Encrypted items: list names if present, or state none were found
- Network artifacts: list names if present, or state none were found
- File extension distribution highlights

=== DATA (facts only) ===
{data_summary_text}
=== END DATA ===""",
            fallback="Detailed findings could not be generated by the AI model. Please refer to the data tables in this report."
        )

        # ── Build PDF ─────────────────────────────────────────────────
        doc = SimpleDocTemplate(pdf_path, pagesize=letter,
                                topMargin=0.55*inch, bottomMargin=0.55*inch,
                                leftMargin=0.6*inch, rightMargin=0.6*inch)

        styles = getSampleStyleSheet()
        # -- Styles --
        cover_title = ParagraphStyle('CoverTitle', parent=styles['Title'],
            fontSize=28, textColor=colors.HexColor('#1a237e'), spaceAfter=6, alignment=1,
            fontName='Helvetica-Bold', leading=34)
        cover_sub = ParagraphStyle('CoverSub', parent=styles['Normal'],
            fontSize=12, textColor=colors.HexColor('#455a64'), alignment=1, spaceAfter=4, leading=16)
        heading_style = ParagraphStyle('SectionHead', parent=styles['Heading2'],
            fontSize=13, textColor=colors.white, spaceAfter=10, spaceBefore=16,
            fontName='Helvetica-Bold', backColor=colors.HexColor('#1a237e'),
            borderPadding=(6, 6, 6, 6), leading=17)
        normal_style = ParagraphStyle('Body', parent=styles['Normal'],
            fontSize=9, textColor=colors.HexColor('#212121'), spaceAfter=5, leading=13.5)
        label_style = ParagraphStyle('Label', parent=styles['Normal'],
            fontSize=8, textColor=colors.HexColor('#616161'), spaceAfter=2, leading=11)

        tbl_style_header = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a237e')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 8),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
            ('TOPPADDING', (0, 0), (-1, -1), 5),
            ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#BDBDBD')),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#E8EAF6')]),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ])

        sus_tbl_style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#b71c1c')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 7.5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#BDBDBD')),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#FFEBEE')]),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ])

        content = []
        section_num = 0  # dynamic section counter

        # ══════════════════════════════════════════════════════════════
        # COVER PAGE
        # ══════════════════════════════════════════════════════════════
        content.append(Spacer(1, 2.0*inch))
        content.append(Paragraph("FORENSIC ANALYSIS REPORT", cover_title))
        content.append(Spacer(1, 0.15*inch))
        # Thin decorative line
        cover_line_data = [['']]
        cover_line_tbl = Table(cover_line_data, colWidths=[5*inch], rowHeights=[2])
        cover_line_tbl.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#1a237e')),
            ('LINEBELOW', (0, 0), (-1, -1), 0, colors.white),
        ]))
        content.append(cover_line_tbl)
        content.append(Spacer(1, 0.3*inch))
        content.append(Paragraph(f"Case ID: {analysis_id}", cover_sub))
        content.append(Paragraph(f"Date Generated: {results['timestamp']}", cover_sub))
        content.append(Paragraph(f"Evidence File: {os.path.basename(image_path)}", cover_sub))
        content.append(Paragraph(f"Filesystem: {data.get('partition_info', {}).get('filesystem_type', 'Unknown')}", cover_sub))
        content.append(Spacer(1, 1.5*inch))
        content.append(Paragraph("CONFIDENTIAL — FOR AUTHORIZED PERSONNEL ONLY",
            ParagraphStyle('Conf', parent=styles['Normal'], fontSize=10,
                           textColor=colors.HexColor('#c62828'), alignment=1, fontName='Helvetica-Bold')))
        content.append(PageBreak())

        # ══════════════════════════════════════════════════════════════
        # SECTION: EVIDENCE INFORMATION
        # ══════════════════════════════════════════════════════════════
        section_num += 1
        content.append(Paragraph(f"{section_num}. EVIDENCE INFORMATION", heading_style))
        case_data = [
            ['Field', 'Value'],
            ['Report Generated', results['timestamp']],
            ['Image File', image_path],
            ['Filesystem Type', data.get('partition_info', {}).get('filesystem_type', 'Unknown')],
            ['Block Size', f"{data.get('partition_info', {}).get('block_size', 'N/A')} bytes"],
            ['Block Count', str(data.get('partition_info', {}).get('block_count', 'N/A'))],
            ['Scan Timestamp', scan_ts],
        ]
        case_table = Table(case_data, colWidths=[1.8*inch, 5.0*inch])
        case_table.setStyle(tbl_style_header)
        content.append(case_table)
        content.append(Spacer(1, 0.2*inch))

        # ══════════════════════════════════════════════════════════════
        # SECTION: EXECUTIVE SUMMARY (LLM)
        # ══════════════════════════════════════════════════════════════
        section_num += 1
        content.append(Paragraph(f"{section_num}. EXECUTIVE SUMMARY", heading_style))
        for para in exec_summary.split('\n\n'):
            para = para.strip()
            if para:
                content.append(Paragraph(para.replace('\n', '<br/>'), normal_style))
        content.append(Spacer(1, 0.15*inch))

        # ══════════════════════════════════════════════════════════════
        # SECTION: ANALYSIS STATISTICS
        # ══════════════════════════════════════════════════════════════
        section_num += 1
        content.append(Paragraph(f"{section_num}. ANALYSIS STATISTICS", heading_style))
        stats_data = [
            ['Metric', 'Value'],
            ['Total Files Scanned', f"{summary.get('total_files', 0):,}"],
            ['Deleted Files Found', str(summary.get('total_deleted_files', 0))],
            ['Suspicious Files (by extension)', str(summary.get('suspicious_files_count', 0))],
            ['Encrypted Items', str(summary.get('encrypted_items_count', 0))],
            ['Network Artifacts', str(summary.get('network_artifacts_count', 0))],
            ['Distinct File Types', str(len(data.get('file_extension_statistics', {})))],
        ]
        stats_table = Table(stats_data, colWidths=[3.0*inch, 3.8*inch])
        stats_table.setStyle(tbl_style_header)
        content.append(stats_table)
        content.append(Spacer(1, 0.2*inch))

        # ══════════════════════════════════════════════════════════════
        # SECTION: FILE EXTENSION DISTRIBUTION
        # ══════════════════════════════════════════════════════════════
        section_num += 1
        content.append(Paragraph(f"{section_num}. FILE EXTENSION DISTRIBUTION", heading_style))
        extensions = data.get('file_extension_statistics', {})
        total_files = max(summary.get('total_files', 1), 1)
        if extensions:
            ext_data = [['Extension', 'Count', 'Percentage']]
            for ext, count in sorted(extensions.items(), key=lambda x: x[1], reverse=True)[:20]:
                pct = f"{(count / total_files) * 100:.1f}%"
                ext_data.append([ext, str(count), pct])
            ext_table = Table(ext_data, colWidths=[2.2*inch, 2.3*inch, 2.3*inch])
            ext_table.setStyle(tbl_style_header)
            content.append(ext_table)
        else:
            content.append(Paragraph("No file extension data available.", normal_style))
        content.append(Spacer(1, 0.2*inch))

        # ══════════════════════════════════════════════════════════════
        # SECTION: SUSPICIOUS FILES
        # ══════════════════════════════════════════════════════════════
        files = data.get('files', {})
        suspicious_files = [
            f for f in files.values()
            if any(f['name'].lower().endswith(ext) for ext in ['.exe', '.bat', '.ps1', '.dll', '.scr', '.vbs'])
        ]
        section_num += 1
        content.append(PageBreak())
        content.append(Paragraph(f"{section_num}. SUSPICIOUS FILES", heading_style))
        if suspicious_files:
            content.append(Paragraph(
                f"<b>{len(suspicious_files)}</b> file(s) with potentially dangerous extensions "
                f"(.exe, .bat, .ps1, .dll, .scr, .vbs) were identified in the disk image.",
                normal_style))
            content.append(Spacer(1, 0.08*inch))
            sus_data = [['#', 'File Name', 'Path', 'Size', 'Modified', 'Deleted']]
            for idx, f in enumerate(sorted(suspicious_files, key=lambda x: x['size'], reverse=True)[:30], 1):
                sus_data.append([
                    str(idx), f['name'][:35], f['path'][:45],
                    _fmt_size(f['size']), f['modification_time'][:16],
                    'YES' if f.get('is_deleted') else 'No'
                ])
            sus_table = Table(sus_data, colWidths=[0.3*inch, 1.3*inch, 1.8*inch, 0.8*inch, 1.2*inch, 0.5*inch])
            sus_table.setStyle(sus_tbl_style)
            content.append(sus_table)
        else:
            content.append(Paragraph(
                "No files with suspicious extensions (.exe, .bat, .ps1, .dll, .scr, .vbs) were found in this disk image.",
                normal_style))
        content.append(Spacer(1, 0.15*inch))

        # ══════════════════════════════════════════════════════════════
        # SECTION: DELETED FILES
        # ══════════════════════════════════════════════════════════════
        deleted_files = [f for f in files.values() if f.get('is_deleted')]
        section_num += 1
        content.append(Paragraph(f"{section_num}. DELETED FILES", heading_style))
        if deleted_files:
            content.append(Paragraph(
                f"<b>{len(deleted_files)}</b> deleted file(s) were recovered from the disk image.", normal_style))
            del_data = [['#', 'File Name', 'Path', 'Size', 'Modified']]
            for idx, f in enumerate(sorted(deleted_files, key=lambda x: x['modification_time'], reverse=True)[:25], 1):
                del_data.append([str(idx), f['name'][:35], f['path'][:50],
                                 _fmt_size(f['size']), f['modification_time'][:16]])
            del_table = Table(del_data, colWidths=[0.3*inch, 1.5*inch, 2.2*inch, 0.9*inch, 1.2*inch])
            del_table.setStyle(tbl_style_header)
            content.append(del_table)
        else:
            content.append(Paragraph("No deleted files were found in this disk image.", normal_style))
        content.append(Spacer(1, 0.15*inch))

        # ══════════════════════════════════════════════════════════════
        # SECTION: ENCRYPTED ITEMS
        # ══════════════════════════════════════════════════════════════
        encrypted = data.get('encrypted_items', [])
        section_num += 1
        content.append(Paragraph(f"{section_num}. ENCRYPTED ITEMS", heading_style))
        if encrypted:
            content.append(Paragraph(
                f"<b>{len(encrypted)}</b> encrypted item(s) were detected.", normal_style))
            enc_data = [['#', 'Name', 'Path', 'Type', 'Size', 'Modified']]
            for idx, item in enumerate(encrypted[:15], 1):
                enc_data.append([
                    str(idx), item['name'][:28], item['path'][:38],
                    item.get('type', 'N/A'), _fmt_size(item['size']),
                    item.get('modified', 'N/A')[:16]])
            enc_table = Table(enc_data, colWidths=[0.3*inch, 1.1*inch, 1.6*inch, 0.8*inch, 0.9*inch, 1.1*inch])
            enc_table.setStyle(tbl_style_header)
            content.append(enc_table)
        else:
            content.append(Paragraph("No encrypted items were found in this disk image.", normal_style))
        content.append(Spacer(1, 0.15*inch))

        # ══════════════════════════════════════════════════════════════
        # SECTION: NETWORK ARTIFACTS
        # ══════════════════════════════════════════════════════════════
        network = data.get('network_artifacts', [])
        section_num += 1
        content.append(Paragraph(f"{section_num}. NETWORK ARTIFACTS", heading_style))
        if network:
            content.append(Paragraph(
                f"<b>{len(network)}</b> network-related artifact(s) were detected.", normal_style))
            net_data = [['#', 'Name', 'Path', 'Type', 'Size', 'Modified']]
            for idx, item in enumerate(network[:15], 1):
                net_data.append([
                    str(idx), item['name'][:28], item['path'][:38],
                    item.get('type', 'N/A'), _fmt_size(item['size']),
                    item.get('modified', 'N/A')[:16]])
            net_table = Table(net_data, colWidths=[0.3*inch, 1.1*inch, 1.6*inch, 0.8*inch, 0.9*inch, 1.1*inch])
            net_table.setStyle(tbl_style_header)
            content.append(net_table)
        else:
            content.append(Paragraph("No network artifacts were found in this disk image.", normal_style))
        content.append(Spacer(1, 0.15*inch))

        # ══════════════════════════════════════════════════════════════
        # SECTION: DETAILED FINDINGS (LLM)
        # ══════════════════════════════════════════════════════════════
        content.append(PageBreak())
        section_num += 1
        content.append(Paragraph(f"{section_num}. DETAILED FINDINGS", heading_style))
        for para in findings_narrative.split('\n\n'):
            para = para.strip()
            if para:
                content.append(Paragraph(para.replace('\n', '<br/>'), normal_style))
        content.append(Spacer(1, 0.2*inch))

        # ══════════════════════════════════════════════════════════════
        # DISCLAIMER
        # ══════════════════════════════════════════════════════════════
        content.append(Spacer(1, 0.4*inch))
        content.append(Paragraph(
            "<i>DISCLAIMER: This report was generated automatically by the Forensic Analysis Suite. "
            "Narrative sections were produced by a local AI model (Llama 3) using only the data "
            "extracted from the uploaded disk image. No external data or assumptions were used. "
            "All findings must be independently verified by a qualified forensic examiner before "
            "being used in any legal or investigative proceedings.</i>",
            ParagraphStyle('Disclaimer', parent=styles['Normal'], fontSize=7,
                           textColor=colors.HexColor('#999999'), alignment=1, leading=10)))

        doc.build(content, onFirstPage=_pdf_page_footer, onLaterPages=_pdf_page_footer)
        return pdf_path

    except Exception as e:
        import traceback; traceback.print_exc()
        print(f"[-] PDF Generation Error: {str(e)}")
        return None

def generate_doc_report(analysis_id, results):
    """Generate a professional Word (DOCX) forensic report with fact-only LLM analysis."""
    try:
        # Lazy import — only loaded when user actually requests DOCX export
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc_path = os.path.join(OUTPUT_FOLDER, f'{analysis_id}_report.docx')
        data = results['data']['forensic_report']
        summary = data['summary']
        data_summary_text = _build_report_data_summary(data)
        image_path = data.get('image_path', 'N/A')

        # ── LLM-generated sections — STRICT fact-only prompts ─────────
        exec_summary = _llm_generate_section(
            f"""You are a digital forensics report writer.

STRICT RULES — you MUST follow every one:
1. ONLY use the facts provided in the DATA section below. Do NOT assume, speculate, or invent any information.
2. If a count is zero, state clearly that none were found. Do NOT fabricate threats or risks.
3. Do NOT reference external tools, external scans, or anything not in the data.
4. Do NOT use markdown formatting. Write plain text only.
5. Write in formal, professional forensic report language.

Write an EXECUTIVE SUMMARY (3-5 short paragraphs) covering ONLY:
- What disk image was analyzed (use the exact image path from the data)
- Filesystem type and scan date from the data
- Exact counts: total files, deleted files, suspicious files, encrypted items, network artifacts
- A factual statement of whether items of interest were found or not

=== DATA (facts only) ===
{data_summary_text}
=== END DATA ===""",
            fallback="Executive summary could not be generated by the AI model."
        )

        findings_narrative = _llm_generate_section(
            f"""You are a digital forensics report writer.

STRICT RULES:
1. ONLY reference facts from the DATA section below. Do NOT assume, speculate, or add information not in the data.
2. If a category has zero items, state that clearly and move on.
3. Do NOT use markdown. Write plain text only.
4. Write in formal forensic report language.

Write a DETAILED FINDINGS section (4-6 paragraphs) covering:
- Suspicious files: list each by name if present, or state none were found
- Deleted files: list names if present, or state none were found
- Encrypted items: list names if present, or state none were found
- Network artifacts: list names if present, or state none were found
- File extension distribution highlights

=== DATA (facts only) ===
{data_summary_text}
=== END DATA ===""",
            fallback="Detailed findings could not be generated by the AI model."
        )

        # ── Build Document ────────────────────────────────────────────
        doc = Document()
        section_num = 0

        # -- Cover Title
        title = doc.add_heading('FORENSIC ANALYSIS REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = RGBColor(26, 35, 126)

        case_para = doc.add_paragraph()
        case_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = case_para.add_run(f'Case ID: {analysis_id}\n')
        run1.font.color.rgb = RGBColor(100, 100, 100)
        run2 = case_para.add_run(f'Generated: {results["timestamp"]}\n')
        run2.font.color.rgb = RGBColor(100, 100, 100)
        run3 = case_para.add_run(f'Evidence File: {os.path.basename(image_path)}')
        run3.font.color.rgb = RGBColor(100, 100, 100)

        conf_para = doc.add_paragraph()
        conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        conf_run = conf_para.add_run('CONFIDENTIAL — FOR AUTHORIZED PERSONNEL ONLY')
        conf_run.font.color.rgb = RGBColor(198, 40, 40)
        conf_run.bold = True
        conf_run.font.size = Pt(11)

        doc.add_page_break()

        # -- Evidence Information
        section_num += 1
        doc.add_heading(f'{section_num}. EVIDENCE INFORMATION', level=1)
        info_rows = [
            ('Image Path', image_path),
            ('Analysis Date', results['timestamp']),
            ('Filesystem Type', data.get('partition_info', {}).get('filesystem_type', 'Unknown')),
            ('Block Size', f"{data.get('partition_info', {}).get('block_size', 'N/A')} bytes"),
            ('Block Count', str(data.get('partition_info', {}).get('block_count', 'N/A'))),
            ('Scan Timestamp', data.get('scan_timestamp', 'N/A')),
        ]
        info_table = doc.add_table(rows=len(info_rows), cols=2)
        info_table.style = 'Light Grid Accent 1'
        for i, (label, value) in enumerate(info_rows):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = str(value)
            for paragraph in info_table.cell(i, 0).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        doc.add_paragraph()

        # -- Executive Summary (LLM)
        section_num += 1
        doc.add_heading(f'{section_num}. EXECUTIVE SUMMARY', level=1)
        for para_text in exec_summary.split('\n\n'):
            para_text = para_text.strip()
            if para_text:
                doc.add_paragraph(para_text)
        doc.add_paragraph()

        # -- Statistics
        section_num += 1
        doc.add_heading(f'{section_num}. ANALYSIS STATISTICS', level=1)
        stats_rows = [
            ('Total Files Scanned', f"{summary.get('total_files', 0):,}"),
            ('Deleted Files Found', str(summary.get('total_deleted_files', 0))),
            ('Suspicious Files (by extension)', str(summary.get('suspicious_files_count', 0))),
            ('Encrypted Items', str(summary.get('encrypted_items_count', 0))),
            ('Network Artifacts', str(summary.get('network_artifacts_count', 0))),
            ('Distinct File Types', str(len(data.get('file_extension_statistics', {})))),
        ]
        stats_table = doc.add_table(rows=len(stats_rows), cols=2)
        stats_table.style = 'Light Grid Accent 1'
        for i, (label, value) in enumerate(stats_rows):
            stats_table.cell(i, 0).text = label
            stats_table.cell(i, 1).text = value
        doc.add_paragraph()

        # -- File Extensions
        section_num += 1
        doc.add_heading(f'{section_num}. FILE EXTENSION DISTRIBUTION', level=1)
        extensions = data.get('file_extension_statistics', {})
        total_files = max(summary.get('total_files', 1), 1)
        sorted_ext = sorted(extensions.items(), key=lambda x: x[1], reverse=True)[:20]
        if sorted_ext:
            ext_table = doc.add_table(rows=len(sorted_ext) + 1, cols=3)
            ext_table.style = 'Light Grid Accent 1'
            ext_table.cell(0, 0).text = 'Extension'
            ext_table.cell(0, 1).text = 'Count'
            ext_table.cell(0, 2).text = 'Percentage'
            for idx, (ext, count) in enumerate(sorted_ext):
                ext_table.cell(idx + 1, 0).text = str(ext)
                ext_table.cell(idx + 1, 1).text = str(count)
                ext_table.cell(idx + 1, 2).text = f"{(count / total_files) * 100:.1f}%"
        else:
            doc.add_paragraph('No file extension data available.')
        doc.add_page_break()

        # -- Suspicious Files
        files = data.get('files', {})
        suspicious_files = [
            f for f in files.values()
            if any(f['name'].lower().endswith(ext) for ext in ['.exe', '.bat', '.ps1', '.dll', '.scr', '.vbs'])
        ]
        section_num += 1
        doc.add_heading(f'{section_num}. SUSPICIOUS FILES', level=1)
        if suspicious_files:
            doc.add_paragraph(
                f'{len(suspicious_files)} file(s) with potentially dangerous extensions '
                f'(.exe, .bat, .ps1, .dll, .scr, .vbs) were identified:')
            sus_table = doc.add_table(rows=min(len(suspicious_files), 30) + 1, cols=6)
            sus_table.style = 'Light Grid Accent 2'
            for i, h in enumerate(['#', 'File Name', 'Path', 'Size', 'Modified', 'Deleted']):
                sus_table.cell(0, i).text = h
            for idx, f in enumerate(sorted(suspicious_files, key=lambda x: x['size'], reverse=True)[:30], 1):
                sus_table.cell(idx, 0).text = str(idx)
                sus_table.cell(idx, 1).text = f['name'][:40]
                sus_table.cell(idx, 2).text = f['path'][:50]
                sus_table.cell(idx, 3).text = _fmt_size(f['size'])
                sus_table.cell(idx, 4).text = f['modification_time'][:16]
                sus_table.cell(idx, 5).text = 'YES' if f.get('is_deleted') else 'No'
        else:
            doc.add_paragraph(
                'No files with suspicious extensions (.exe, .bat, .ps1, .dll, .scr, .vbs) '
                'were found in this disk image.')
        doc.add_paragraph()

        # -- Deleted Files
        deleted_files = [f for f in files.values() if f.get('is_deleted')]
        section_num += 1
        doc.add_heading(f'{section_num}. DELETED FILES', level=1)
        if deleted_files:
            doc.add_paragraph(f'{len(deleted_files)} deleted file(s) were recovered from the disk image:')
            del_table = doc.add_table(rows=min(len(deleted_files), 25) + 1, cols=5)
            del_table.style = 'Light Grid Accent 1'
            for i, h in enumerate(['#', 'File Name', 'Path', 'Size', 'Modified']):
                del_table.cell(0, i).text = h
            for idx, f in enumerate(sorted(deleted_files, key=lambda x: x['modification_time'], reverse=True)[:25], 1):
                del_table.cell(idx, 0).text = str(idx)
                del_table.cell(idx, 1).text = f['name'][:40]
                del_table.cell(idx, 2).text = f['path'][:50]
                del_table.cell(idx, 3).text = _fmt_size(f['size'])
                del_table.cell(idx, 4).text = f['modification_time'][:16]
        else:
            doc.add_paragraph('No deleted files were found in this disk image.')
        doc.add_paragraph()

        # -- Encrypted Items
        encrypted = data.get('encrypted_items', [])
        section_num += 1
        doc.add_heading(f'{section_num}. ENCRYPTED ITEMS', level=1)
        if encrypted:
            doc.add_paragraph(f'{len(encrypted)} encrypted item(s) were detected:')
            enc_table = doc.add_table(rows=min(len(encrypted), 15) + 1, cols=5)
            enc_table.style = 'Light Grid Accent 1'
            for i, h in enumerate(['#', 'Name', 'Path', 'Type', 'Size']):
                enc_table.cell(0, i).text = h
            for idx, item in enumerate(encrypted[:15], 1):
                enc_table.cell(idx, 0).text = str(idx)
                enc_table.cell(idx, 1).text = item['name'][:35]
                enc_table.cell(idx, 2).text = item['path'][:45]
                enc_table.cell(idx, 3).text = item.get('type', 'N/A')
                enc_table.cell(idx, 4).text = _fmt_size(item['size'])
        else:
            doc.add_paragraph('No encrypted items were found in this disk image.')

        # -- Network Artifacts
        network = data.get('network_artifacts', [])
        section_num += 1
        doc.add_heading(f'{section_num}. NETWORK ARTIFACTS', level=1)
        if network:
            doc.add_paragraph(f'{len(network)} network-related artifact(s) were detected:')
            net_table = doc.add_table(rows=min(len(network), 15) + 1, cols=5)
            net_table.style = 'Light Grid Accent 1'
            for i, h in enumerate(['#', 'Name', 'Path', 'Type', 'Size']):
                net_table.cell(0, i).text = h
            for idx, item in enumerate(network[:15], 1):
                net_table.cell(idx, 0).text = str(idx)
                net_table.cell(idx, 1).text = item['name'][:35]
                net_table.cell(idx, 2).text = item['path'][:45]
                net_table.cell(idx, 3).text = item.get('type', 'N/A')
                net_table.cell(idx, 4).text = _fmt_size(item['size'])
        else:
            doc.add_paragraph('No network artifacts were found in this disk image.')

        # -- Detailed Findings (LLM)
        doc.add_page_break()
        section_num += 1
        doc.add_heading(f'{section_num}. DETAILED FINDINGS', level=1)
        for para_text in findings_narrative.split('\n\n'):
            para_text = para_text.strip()
            if para_text:
                doc.add_paragraph(para_text)

        # -- Disclaimer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer_run = footer.add_run(
            'DISCLAIMER: This report was generated automatically by the Forensic Analysis Suite. '
            'Narrative sections were produced by a local AI model (Llama 3) using only the data '
            'extracted from the uploaded disk image. No external data or assumptions were used. '
            'All findings must be independently verified by a qualified forensic examiner before '
            'being used in any legal or investigative proceedings.'
        )
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(150, 150, 150)
        footer_run.italic = True

        doc.save(doc_path)
        return doc_path

    except Exception as e:
        import traceback; traceback.print_exc()
        print(f"[-] DOC Generation Error: {str(e)}")
        return None

# ============================================================================
# LAZY-LOAD RAG ENGINE (only when needed to avoid slow startup)
# ============================================================================
_rag_loaded = False
_vectorize_and_store = None
_rag_ask = None

def _ensure_rag_loaded():
    """Lazy-load RAG engine on first use."""
    global _rag_loaded, _vectorize_and_store, _rag_ask
    if not _rag_loaded:
        print("[*] Loading RAG engine...")
        from rag_engine import vectorize_and_store, ask as rag_ask
        _vectorize_and_store = vectorize_and_store
        _rag_ask = rag_ask
        _rag_loaded = True
        print("[+] RAG engine loaded.")

# ============================================================================
# FLASK ROUTES
# ============================================================================

@app.route('/')
def index():
    """Render main dashboard."""
    return render_template('index.html')

@app.route('/api/status')
def get_status():
    """Get current analysis status."""
    return jsonify(analysis_status)

@app.route('/api/upload', methods=['POST'])
def upload_file():
    """Handle file upload."""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file type. Allowed: raw, dd, e01, img, iso, bin, dmg'}), 400
        
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{timestamp}_{filename}"
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        file.save(filepath)
        
        return jsonify({
            'success': True,
            'message': 'File uploaded successfully',
            'filename': filename,
            'filepath': filepath
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/analysis/<analysis_id>/status', methods=['GET'])
def get_analysis_status(analysis_id):
    """Get status of a specific analysis."""
    global analysis_status
    
    try:
        # Check if this is the current analysis
        if analysis_status['current'] == analysis_id:
            return jsonify({
                'id': analysis_id,
                'status': analysis_status['status'],
                'progress': analysis_status['progress'],
                'message': analysis_status['message'],
                'errors': analysis_status['errors'],
                'results': analysis_status['results']
            })
        else:
            return jsonify({'error': 'Analysis not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/analyze', methods=['POST'])
def analyze():
    """Start forensic analysis."""
    global analysis_status
    
    try:
        data = request.json
        image_path = data.get('image_path')
        
        if not image_path:
            return jsonify({'error': 'No image path provided'}), 400
        
        if not os.path.exists(image_path):
            return jsonify({'error': f'Image file not found: {image_path}'}), 400
        
        analysis_id = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        # Start analysis in background thread
        analysis_thread = threading.Thread(
            target=run_forensic_analysis,
            args=(image_path, analysis_id)
        )
        analysis_thread.daemon = True
        analysis_thread.start()
        
        return jsonify({
            'success': True,
            'analysis_id': analysis_id,
            'message': 'Analysis started'
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/export/<format>', methods=['GET'])
def export_report(format):
    """Export report as PDF or DOC."""
    global analysis_status
    
    try:
        if not analysis_status['results']:
            return jsonify({'error': 'No analysis results available'}), 400
        
        results = analysis_status['results']
        analysis_id = results['analysis_id']
        
        if format.lower() == 'pdf':
            pdf_path = generate_pdf_report(analysis_id, results)
            if pdf_path and os.path.exists(pdf_path):
                return send_file(pdf_path, mimetype='application/pdf', 
                               as_attachment=True, 
                               download_name=f'forensic_report_{analysis_id}.pdf')
            else:
                return jsonify({'error': 'Failed to generate PDF'}), 500
        
        elif format.lower() == 'doc' or format.lower() == 'docx':
            doc_path = generate_doc_report(analysis_id, results)
            if doc_path and os.path.exists(doc_path):
                return send_file(doc_path, 
                               mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                               as_attachment=True,
                               download_name=f'forensic_report_{analysis_id}.docx')
            else:
                return jsonify({'error': 'Failed to generate DOCX'}), 500
        
        else:
            return jsonify({'error': 'Invalid format. Use "pdf" or "doc"'}), 400
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/results')
def get_results():
    """Get analysis results."""
    global analysis_status
    
    if not analysis_status['results']:
        return jsonify({'error': 'No results available'}), 400
    
    # Return results summary (don't send entire JSON to avoid size issues)
    results = analysis_status['results']
    data = results['data']['forensic_report']
    
    return jsonify({
        'timestamp': results['timestamp'],
        'image_path': results['image_path'],
        'summary': data['summary'],
        'extensions': data.get('file_extension_statistics', {}),
        'encrypted_items': data.get('encrypted_items', []),
        'network_artifacts': data.get('network_artifacts', [])
    })

@app.route('/api/reset')
def reset_analysis():
    """Reset analysis for new one."""
    global analysis_status
    
    analysis_status = {
        'current': None,
        'progress': 0,
        'status': 'idle',
        'message': '',
        'errors': [],
        'results': None,
        'rag_collection': None
    }
    
    return jsonify({'success': True, 'message': 'Analysis reset'})


# ============================================================================
# RAG CHATBOT ENDPOINTS
# ============================================================================

@app.route('/api/rag/chat', methods=['POST'])
def rag_chat():
    """Answer a user question using RAG over the vectorized forensic report."""
    global analysis_status
    try:
        _ensure_rag_loaded()  # Lazy-load RAG on first use
        
        data = request.json
        query = data.get('query', '').strip()
        if not query:
            return jsonify({'error': 'No query provided'}), 400

        collection_name = (
            data.get('collection')
            or (analysis_status.get('results') or {}).get('rag_collection')
            or analysis_status.get('rag_collection')
        )
        if not collection_name:
            return jsonify({'error': 'No forensic report has been vectorized yet. Please upload and analyze a file first.'}), 400

        top_k = int(data.get('top_k', 5))
        use_llm = bool(data.get('use_llm', True))

        result = _rag_ask(collection_name, query, top_k=top_k, use_llm=use_llm)
        return jsonify({
            'success': True,
            'answer': result['answer'],
            'context': result['context'],
            'collection': result['collection'],
            'rejected': result.get('rejected', False),
        })
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/rag/status', methods=['GET'])
def rag_status():
    """Check if RAG is ready (collection exists)."""
    global analysis_status
    collection_name = (
        (analysis_status.get('results') or {}).get('rag_collection')
        or analysis_status.get('rag_collection')
    )
    return jsonify({
        'ready': collection_name is not None,
        'collection': collection_name,
    })


def _warmup_models():
    """Pre-load the embedding model in a background thread so the first
    analysis / RAG query doesn't pay a cold-start penalty."""
    try:
        _ensure_rag_loaded()
        from rag_engine import _get_embed_model
        _get_embed_model()
        print("[+] Embedding model pre-warmed successfully.")
    except Exception as e:
        print(f"[!] Model warmup warning (non-fatal): {e}")


if __name__ == '__main__':
    print("\n" + "="*80)
    print(" FORENSIC ANALYSIS WEB APPLICATION - STARTING")
    print("="*80)
    print("\n[+] Frontend: http://localhost:5001")
    print("[+] Backend: Flask Server")
    print("[+] Upload Folder: " + UPLOAD_FOLDER)
    print("[+] Output Folder: " + OUTPUT_FOLDER)
    print("\n" + "="*80 + "\n")

    # Pre-warm the embedding model on a background thread so it's ready
    # before the first upload/RAG query (avoids ~3-5s cold-start delay)
    warmup_thread = threading.Thread(target=_warmup_models, daemon=True)
    warmup_thread.start()
    
    app.run(debug=False, host='127.0.0.1', port=5001, threaded=True)
